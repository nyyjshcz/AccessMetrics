"""Build blank, explicitly conditional acceptance templates.

The files intentionally contain no names, dates, signatures, stamps, or claims
that a presentation/submission/acceptance already happened. They are generated
from the repository Markdown templates so the DOCX and PDF wording stays in
lockstep.
"""
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "acceptance-materials"
WATERMARK = "草稿 / 未签署 / 仅供接收单位审核，不能作为项目已展示、已提交或已接收的证明。"


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 8),
        ("Heading 1", 15, "2E74B5", 14, 7),
        ("Heading 2", 12.5, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AccessCheck Lishui · 条件式成果接收模板")
    run.font.name = "SimSun"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(WATERMARK)
    run.font.name = "SimSun"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(155, 28, 28)


def qn(tag: str):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "SimSun"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(122, 90, 0)
    shading = cell._tc.get_or_add_tcPr().append(
        __import__("docx").oxml.parse_xml(
            '<w:shd {} w:fill="FFF8E8"/>'.format(
                __import__("docx").oxml.ns.nsdecls("w")
            )
        )
    )
    del shading
    doc.add_paragraph()


def add_fields(doc: Document, fields: list[str]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for field in fields:
        cells = table.add_row().cells
        cells[0].width = Inches(1.7)
        cells[1].width = Inches(4.5)
        cells[0].text = field
        cells[1].text = "（由负责人或接收单位按真实情况填写）"
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "SimSun"
                    run.font.size = Pt(10)
        cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def make_docx(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]], fields: list[str]) -> None:
    doc = Document()
    set_doc_defaults(doc)
    doc.add_paragraph(title, style="Title")
    p = doc.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    add_callout(doc, WATERMARK)
    doc.add_heading("使用边界", level=1)
    doc.add_paragraph(
        "本文件是空白条件式模板。只有在真实展示、提交或接收发生后，才可由相应负责人填写事实字段；AI 不代填姓名、日期、单位意见、签字或盖章。"
    )
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)
    doc.add_heading("待填写事实", level=1)
    add_fields(doc, fields)
    doc.add_heading("确认", level=1)
    doc.add_paragraph(
        "填写完成后，应由负责人核对事实并将签署件/盖章件放入 Git 外的 private-inputs/signed-acceptance/；仓库内只保留本空白模板。"
    )
    doc.save(path)


def register_cjk_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("AccessCJK", str(candidate)))
            return "AccessCJK"
    pdfmetrics.registerFont(pdfmetrics.getFont("STSong-Light"))
    return "STSong-Light"


def make_pdf(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]], fields: list[str]) -> None:
    font = register_cjk_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("AccessTitle", parent=styles["Title"], fontName=font, fontSize=20, leading=26, textColor=colors.HexColor("#0B2545"), alignment=TA_LEFT, spaceAfter=7)
    subtitle_style = ParagraphStyle("AccessSubtitle", parent=styles["Normal"], fontName=font, fontSize=10, leading=15, textColor=colors.HexColor("#555555"), spaceAfter=12)
    h1 = ParagraphStyle("AccessH1", parent=styles["Heading1"], fontName=font, fontSize=14, leading=20, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("AccessBody", parent=styles["BodyText"], fontName=font, fontSize=10.5, leading=16, spaceAfter=7)
    note = ParagraphStyle("AccessNote", parent=body, fontName=font, textColor=colors.HexColor("#7A5A00"), backColor=colors.HexColor("#FFF8E8"), borderColor=colors.HexColor("#E7D69B"), borderWidth=0.6, borderPadding=7, spaceAfter=12)
    story = [Paragraph(title, title_style), Paragraph(subtitle, subtitle_style), Paragraph(WATERMARK, note)]
    story += [Paragraph("使用边界", h1), Paragraph("本文件是空白条件式模板。只有在真实展示、提交或接收发生后，才可由相应负责人填写事实字段；AI 不代填姓名、日期、单位意见、签字或盖章。", body)]
    for heading, paragraphs in sections:
        story.append(Paragraph(heading, h1))
        story.extend(Paragraph(text, body) for text in paragraphs)
    story.append(Paragraph("待填写事实", h1))
    table_data = [[field, "（由负责人或接收单位按真实情况填写）"] for field in fields]
    table = Table(table_data, colWidths=[47 * mm, 113 * mm], repeatRows=0)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("LEADING", (0, 0), (-1, -1), 14),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C9D2DC")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([table, Spacer(1, 8), Paragraph("确认", h1), Paragraph("填写完成后，应由负责人核对事实并将签署件/盖章件放入 Git 外的 private-inputs/signed-acceptance/；仓库内只保留本空白模板。", body)])

    def decorate(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.setFillColor(colors.HexColor("#9B1C1C"))
        canvas.drawCentredString(A4[0] / 2, 12 * mm, WATERMARK)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawRightString(A4[0] - 18 * mm, 12 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=20 * mm, title=title, author="AccessCheck Lishui").build(story, onFirstPage=decorate, onLaterPages=decorate)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    acceptance_sections = [
        ("条件式说明", ["如经现场展示并确认接收，则由接收单位填写以下内容。空白状态不代表已经展示、提交或接收。"]),
        ("项目与材料", ["项目名称、版本、报告范围、导出标识、验证记录和交接材料均应以真实冻结文件为准。"]),
    ]
    acceptance_fields = ["项目名称", "负责人姓名", "接收单位", "展示日期", "提交日期", "接收状态", "接收单位意见", "负责人签字", "接收单位签字/盖章", "填写日期"]
    report_sections = [
        ("报告交付说明", ["本报告仅评价 axe-core 能够自动判断的网页无障碍检查项，不等同于完整人工审计、官方 WCAG 合规认证或“符合 WCAG 的百分比”。"]),
        ("真实数据绑定", ["正式网站、分数、排名、人工复核、敏感性分析和结论只能从通过 R1–R5 的冻结导出生成；没有真实输入时不得填写。"]),
    ]
    report_fields = ["研究协议版本", "正式网站清单", "source export-id", "final export-id", "manifest hash", "报告核对人", "核对日期", "接收单位意见", "签字/盖章"]
    jobs = [
        ("成果接收页", "成果材料核对与接收状态记录", acceptance_sections, acceptance_fields),
        ("项目实践与成果接收证明", "条件式证明模板（未签署）", report_sections, report_fields),
    ]
    for stem, subtitle, sections, fields in jobs:
        make_docx(OUTPUT / f"{stem}.docx", stem, subtitle, sections, fields)
        make_pdf(OUTPUT / f"{stem}.pdf", stem, subtitle, sections, fields)


if __name__ == "__main__":
    main()
